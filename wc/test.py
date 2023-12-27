# import static.txt.pass2txt as pass2txt
# import wc.count_noun as count_noun
# from nltk.corpus import stopwords
# import nltk
# import wc.make_wc as make_wc
# import wc.pass2times as pass2times

from docx import Document
import pathlib
import MeCab
import docx
from wordcloud import WordCloud
import json
import re

n = 10

import sys

def json_sort(path):
    with open(path, "r", encoding="utf-8") as f:
        tfdict = json.load(f)

    sorted_tfdict = sorted(tfdict.items(), key=lambda x: x[1], reverse=True)
    sum = 0
    for val in tfdict.values():
        sum += val
    print(sum)

    with open(f"sorted_{path}", "w", encoding="utf-8") as f:
        json.dump(sorted_tfdict, f, ensure_ascii=False)
        
json_sort("tfdict.json")

def make_wordcloud(dict, html_name):
    wc = WordCloud(
        font_path=r"C:\Users\nutta\myProject\FileSystem\YUMIN.TTF",
        prefer_horizontal=1,
        background_color="white",
        # ~~~~~ここまでデフォルト~~~~~
        # grayでモノクロ、jetでRGB
        colormap="Greys",
        width=600,
        height=600,
        font_step=20,
        min_font_size=20,
        max_font_size=200,
    )
    wc.generate_from_frequencies(dict)

    wc.to_file("wc.png")
    wordcloud = wc.to_svg()

    javascript = """

    <script>
    var textElements = document.getElementsByTagName("text");

    for (var i = 0; i < textElements.length; i++) {
        textElements[i].style.fill = "rgb(0, 0, 0)";
    }
    </script>

    """

    html_name = r"html_wc/" + html_name
    with open(html_name, "w", encoding="utf-8") as f:
        f.write(wordcloud)
        f.write(javascript)


def count_tfidf():
    with open("transposed_tfidf_data.json", "r", encoding="utf-8") as f:
        tfidf = json.load(f)

    new_tfidf = {}

    for file in tfidf.keys():
        tfidf[file] = sorted(tfidf[file].items(), key=lambda x: x[1], reverse=True)
        n = 1
        for word in tfidf[file][:n]:
            if word[0] in new_tfidf:
                new_tfidf[word[0]] += 1
            else:
                new_tfidf[word[0]] = 1

    with open(f"tfidf{n}.json", "w", encoding="utf-8") as f:
        json.dump(new_tfidf, f, ensure_ascii=False)

    return new_tfidf


# count_tfidf()

# new_dict = count_tfidf()

# make_wordcloud(new_dict, "tfidf10.html")

with open("word_index.json", "r", encoding="utf-8") as f:
    words_index = json.load(f)

with open("tfidf10.json", "r", encoding="utf-8") as f:
    tfidf10 = json.load(f)

    new_dfdict = {}

    hinshis = [
        "普通名詞",
        # "固有名詞",
        # "数詞"
    ]

for word in tfidf10.keys():
    for hinshi in hinshis:
        for index_word in words_index[hinshi]:
            if word == index_word:
                new_dfdict[word] = tfidf10[word]
            else:
                continue

# with open("tfidf10_proper.json", "w", encoding="utf-8") as f:
#     json.dump(new_dfdict, f, ensure_ascii=False)

# make_wordcloud(new_dfdict, "tfidf10_common.html")


def hinshi_sep(json_path):
    with open("word_index.json", "r", encoding="utf-8") as f:
        words_index = json.load(f)

    with open(json_path, "r", encoding="utf-8") as f:
        dfdict = json.load(f)

    new_dfdict = {}

    hinshis = [
        # "普通名詞",
        "固有名詞",
        # "数詞"
    ]

    for hinshi in hinshis:
        for word in words_index[hinshi]:
            if word not in new_dfdict:
                new_dfdict[word] = dfdict[word]
            else:
                continue

    return new_dfdict


list = ["tfidf10_common.json", "tfidf10_proper.json", "tfidf10_numeral.json"]
words = " ".join(list)
list2 = ["tfidf10_common.html", "tfidf10_proper.html", "tfidf10_numeral.html"]

words += " ".join(list2)
print(words)

# test_dict = {}
# test_dict["a"] = []
# test_dict["a"].append("b")
# print(test_dict)

# path = "C:\\Users\\nutta\\OneDrive\\ドキュメント\\授業資料"


# def PathtoTxt(path):
#     if ".docx" in str(path):
#         doc = docx.Document(path)
#         words = ""
#         for par in doc.paragraphs:
#             words += par.text.replace("\u2028", "")
#     elif ".txt" in str(path):
#         # 行ごとに分割
#         with open(path, encoding="utf-8") as f:
#             contents = f.readlines()
#         words = "".join(contents)
#     return words


# def checknountype(file, word_cnt、wordlist):
#     words = PathtoTxt(file)
#     tagger = MeCab.Tagger()
#     node = tagger.parseToNode(words).next
#     while node:
#         words = node.feature.split(",")
#         # 表層形,品詞,品詞細分類1,品詞細分類2,品詞細分類3,活用型,活用形,原形,読み,発音
#         # if words[2] in word_cnt:
#         #     word_cnt[words[2]] += 1
#         #     if words[2] == "サ変形状詞可能":
#         #         # print(node.surface, end=" ")
#         #         wordlist.append(node.surface)
#         # elif words[1] == "普通名詞":
#         #     word_cnt[words[2]] = 1

#         if words[0] in word_cnt:
#             word_cnt[words[1]] += 1
#         elif words[0] == "名詞":
#             word_cnt[words[2]] = 1
#         else:
#             pass
#         node = node.next
#     return (word_cnt,wordlist)


# def getfiles(path):
#     p = pathlib.Path(path)
#     docs_list = list(p.glob("*.docx"))
#     wordlist = []
#     for file in docs_list:
#         (word_cnt,wordlist) = checknountype(file, wordlist)

#     for key, value in sorted(word_cnt.items(), key=lambda x: x[1], reverse=True):
#         print(f"{key}:\t {value}")

#     wc = WordCloud(
#         font_path="C:\Windows\Fonts\yumin.ttf",
#         prefer_horizontal=1,
#         background_color="white",
#         # ~~~~~ここまでデフォルト~~~~~
#         # grayでモノクロ、jetでRGB
#         colormap=para.color,
#         width=para.width,
#         height=para.height,
#         font_step=para.step,
#         min_font_size=para.min,
#         max_font_size=para.max,
#     )
#     wc.generate_from_text(wordlist)
#     wc.to_file("wc.png")
#     return docs_list


# getfiles(path)


# dict = {"title":1, "情報":2}

# print(dict.items())

# .docxファイルを開く

# p = pathlib.Path(r"C:\Users\nutta\OneDrive\ドキュメント\授業資料")
# docs_list = list(p.glob("*.docx"))
# p = r'C:\Users\nutta\OneDrive\ドキュメント\授業資料\(学類長・専門学群長宛)平成31年度座長団選出報告書.docx'


# with open("comp_time.txt", "w", encoding="utf-8") as f:
#     for file in docs_list:
#         doc = Document(file)  # ファイル名を実際のファイル名に置き換える
#         title = str(file).split("\\")[-1]

#         # メタデータを取得
# metadata = {
#     # 'Title':title,
#     # 'Title': doc.core_properties.title,
#     # 'Author': doc.core_properties.author,
#     # 'Subject': doc.core_properties.subject,
#     # 'Keywords': doc.core_properties.keywords,
#     # 'Comments': doc.core_properties.comments,
#     # 'Last Modified By': doc.core_properties.last_modified_by,
#     'os_time': pass2times.get_ts(file),
#     'Revision': doc.core_properties.revision,
#     'Created': doc.core_properties.created,
#     'Modified': doc.core_properties.modified,
# }

#         # メタデータを表示
#         for key, value in metadata.items():
#             print(f'{key}: {value}', file = f)
#         print("-" * 20, file = f)

# make_wc.get({"a":10})

# with open("stopword.txt", "r",encoding="utf-8") as f:
#   stopwords = f.readlines()

# stopwords = [string.strip() for string in stopwords if string.strip()]

# for word in stopwords:
#   print(word)

# nltk.download('stopwords')
# stop_words = stopwords.words('english')
# print(stop_words)

# stoppath = "Japanese.txt"
# with open(stoppath, "r", encoding= "utf-8") as f:
#   stopwords = f.readlines()

# # nltkモジュールから取得したstopwords
# with open("English.txt", "r", encoding="utf-8") as f:
#   en_stopwords = f.readlines()

# en_sw = en_stopwords[0].replace(",|[|]| ", "")
# for word in en_sw.split(","):
#   stopwords.append(word[2:-1])

# stopwords = [string.strip() for string in stopwords if string.strip()]

# print(stopwords)

# p = r"C:\Users\nutta\OneDrive\ドキュメント\授業資料\(学類長・専門学群長宛)平成31年度座長団選出報告書.docx"

# p2 = r"C:\Users\nutta\myProject\FileSystem\sample1.txt"

# words = pass2txt.getwords(p2)

# print(count_noun.getcount(words))


# stoppath = "Japanese.txt"
# with open(stoppath, "r", encoding = "utf-8") as f:
#   stopwords = f.readlines()

# stopwords = [string.strip() for string in stopwords if string.strip()]
# print(stopwords[:3])


# svg = document.getElementsByTagName("svg")[0];
# text_tags =  svg.getElementsByTagName("text")
# for(var i=0; i<text_tags.length; i++){
#     text_tags[i].addEventListener(
#         "click",
#         function(){
#             word = this.textContent;
#             console.log(this);
#             this.setAttribute("style", "border: 3px solid;");
#             word_uri = encodeURI(word);
#             url = "https://www.google.com/search?q=" + word_uri;
#             window.open(url, "_bkank");

#         }
#     )
# }
